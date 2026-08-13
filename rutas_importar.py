"""Importacion de contactos y empresas desde archivos, con matching automatico."""
import io

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form
from sqlalchemy import text
from sqlalchemy.orm import Session

from database import obtener_sesion
from autenticacion import usuario_actual

router = APIRouter(prefix="/importar", tags=["Importar"])

EXTENSIONES_TEXTO = (".txt", ".csv", ".md", ".tsv")
EXTENSIONES_WORD = (".docx",)
EXTENSIONES_EXCEL = (".xlsx", ".xlsm")


def leer_word(contenido: bytes) -> str:
    """Extrae el texto de un archivo Word (.docx), incluyendo tablas."""
    from docx import Document

    doc = Document(io.BytesIO(contenido))
    partes = []
    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            partes.append(parrafo.text)
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def leer_excel(contenido: bytes) -> str:
    """Extrae el contenido de un Excel (.xlsx) como texto tabulado."""
    from openpyxl import load_workbook

    libro = load_workbook(io.BytesIO(contenido), data_only=True, read_only=True)
    partes = []
    for hoja in libro.worksheets:
        partes.append(f"=== Hoja: {hoja.title} ===")
        for fila in hoja.iter_rows(values_only=True):
            celdas = ["" if v is None else str(v).strip() for v in fila]
            if any(celdas):
                partes.append(" | ".join(celdas))
    libro.close()
    return "\n".join(partes)


def resolver_coincidencias(db: Session, uid: str, datos: dict, org_forzada: str | None):
    """Asocia automaticamente cada registro con lo que ya existe en la base.

    Tolera diferencias de escritura: apostrofes, tildes, mayusculas, puntos
    y sufijos societarios. No pregunta nada: resuelve solo.
    """
    from coincidencias import buscar_coincidencia

    contactos_bd = [
        {"id": str(f.id), "nombre": f.nombre, "email": f.email}
        for f in db.execute(
            text("""
                SELECT c.id, c.nombre, MIN(ce.email) AS email
                FROM contactos c
                LEFT JOIN contacto_emails ce ON ce.contacto_id = c.id
                WHERE c.usuario_id = :uid
                GROUP BY c.id, c.nombre
            """),
            {"uid": uid},
        ).fetchall()
    ]
    empresas_bd = [
        {"id": str(f.id), "nombre": f.nombre}
        for f in db.execute(
            text("SELECT id, nombre FROM empresas WHERE usuario_id = :uid"),
            {"uid": uid},
        ).fetchall()
    ]
    orgs_bd = [
        {"id": str(f.id), "nombre": f.nombre}
        for f in db.execute(
            text("SELECT id, nombre FROM usuario_organizaciones WHERE usuario_id = :uid"),
            {"uid": uid},
        ).fetchall()
    ]

    # --- Organizacion: la del filtro manda; si no, la detectada en el archivo ---
    org_id, org_nombre = None, None
    if org_forzada and org_forzada not in ("todas", "personal"):
        org_id = org_forzada
        encontrada = next((o for o in orgs_bd if o["id"] == org_forzada), None)
        org_nombre = encontrada["nombre"] if encontrada else None
    elif datos.get("organizacion_detectada"):
        o, _ = buscar_coincidencia(
            datos["organizacion_detectada"], orgs_bd, es_empresa=True, umbral=0.7
        )
        if o:
            org_id, org_nombre = o["id"], o["nombre"]

    # --- Empresas ---
    mapa_empresas = {}
    for e in datos.get("empresas", []):
        nombre = (e.get("nombre") or "").strip()
        if not nombre:
            continue
        existente, puntaje = buscar_coincidencia(
            nombre, empresas_bd, es_empresa=True, umbral=0.75
        )
        if existente:
            e["id_existente"] = existente["id"]
            e["nombre_existente"] = existente["nombre"]
            e["accion"] = "reutilizar"
            mapa_empresas[nombre.lower()] = existente["id"]
        else:
            e["accion"] = "crear"
        e["organizacion_id"] = org_id

    # --- Contactos ---
    for c in datos.get("contactos", []):
        nombre = (c.get("nombre") or "").strip()
        if not nombre:
            continue
        existente = None
        email = (c.get("email") or "").strip().lower()
        if email:
            existente = next(
                (x for x in contactos_bd if (x.get("email") or "").lower() == email), None
            )
        if existente is None:
            existente, _ = buscar_coincidencia(nombre, contactos_bd, umbral=0.8)
        if existente:
            c["id_existente"] = existente["id"]
            c["nombre_existente"] = existente["nombre"]
            c["accion"] = "actualizar"
        else:
            c["accion"] = "crear"
        c["organizacion_id"] = org_id

        emp = (c.get("empresa_nombre") or "").strip()
        if emp:
            if emp.lower() in mapa_empresas:
                c["empresa_id_existente"] = mapa_empresas[emp.lower()]
            else:
                coincide, _ = buscar_coincidencia(
                    emp, empresas_bd, es_empresa=True, umbral=0.75
                )
                if coincide:
                    c["empresa_id_existente"] = coincide["id"]

    # --- Eventos ---
    for ev in datos.get("eventos", []):
        ev["organizacion_id"] = org_id
        nombre_c = (ev.get("contacto_nombre") or "").strip()
        if nombre_c:
            coincide, _ = buscar_coincidencia(nombre_c, contactos_bd, umbral=0.8)
            if coincide:
                ev["contacto_id_existente"] = coincide["id"]

    datos["organizacion_id"] = org_id
    datos["organizacion_nombre"] = org_nombre
    return datos


@router.post("/analizar")
async def analizar_archivo(
    archivo: UploadFile = File(...),
    org: str = Form("todas"),
    db: Session = Depends(obtener_sesion),
    uid: str = Depends(usuario_actual),
):
    """Lee un archivo, extrae datos con Claude y los asocia automaticamente
    con los contactos, empresas y organizaciones que ya existen."""
    from generador_ia import extraer_entidades

    nombre = (archivo.filename or "").lower()
    contenido = await archivo.read()

    if len(contenido) > 5_000_000:
        raise HTTPException(status_code=400, detail="Archivo demasiado grande (max 5 MB)")

    if nombre.endswith(EXTENSIONES_TEXTO):
        try:
            texto = contenido.decode("utf-8")
        except UnicodeDecodeError:
            texto = contenido.decode("latin-1", errors="ignore")
    elif nombre.endswith(EXTENSIONES_WORD):
        try:
            texto = leer_word(contenido)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"No se pudo leer el Word: {e}")
    elif nombre.endswith(EXTENSIONES_EXCEL):
        try:
            texto = leer_excel(contenido)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"No se pudo leer el Excel: {e}")
    elif nombre.endswith(".doc") or nombre.endswith(".xls"):
        raise HTTPException(
            status_code=400,
            detail="Formato antiguo no soportado: guarda el archivo como .docx o .xlsx",
        )
    else:
        raise HTTPException(
            status_code=400,
            detail="Formatos aceptados: .txt, .csv, .md, .tsv, .docx, .xlsx",
        )

    if not texto.strip():
        raise HTTPException(status_code=400, detail="El archivo esta vacio o no contiene texto legible")

    try:
        resultado = extraer_entidades(texto)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error al analizar con IA: {e}")

    return resolver_coincidencias(db, uid, resultado, org)